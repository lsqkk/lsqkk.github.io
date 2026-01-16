import os
import re
from docx import Document as DocxDocument

word_file_path = "blog/daily/2.docx"

if not os.path.exists(word_file_path):
    print(f"错误：找不到文件 '{word_file_path}'")
    exit()

try:
    doc = DocxDocument(word_file_path)
    print(f"文档 '{word_file_path}' 加载成功。")
except Exception as e:
    print(f"加载文档时出现错误: {e}")
    exit()

base_name = os.path.splitext(word_file_path)[0]
txt_file_path = base_name + "_超链接.txt"

hyperlink_list = []

# 方法1：通过解析XML来查找超链接
print("正在解析文档结构...")

# 获取文档的核心XML
doc_element = doc.element.body

# 定义递归函数来查找超链接
def find_hyperlinks(element, current_text=""):
    hyperlinks = []
    
    # 检查当前元素是否是超链接
    if element.tag.endswith('}hyperlink') or element.tag.endswith('}hlink'):
        # 提取链接地址
        r_id = element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        link_url = ""
        
        # 从文档关系中获取实际URL
        if r_id and hasattr(doc.part, 'rels'):
            rel = doc.part.rels.get(r_id)
            if rel:
                link_url = rel.target_ref if hasattr(rel, 'target_ref') else str(rel.target)
        
        # 提取链接文本
        link_text = ""
        for child in element.iter():
            if child.tag.endswith('}t'):  # 文本元素
                if child.text:
                    link_text += child.text
        
        if link_url:
            hyperlinks.append((link_text.strip(), link_url))
    
    # 递归检查子元素
    for child in element:
        hyperlinks.extend(find_hyperlinks(child))
    
    return hyperlinks

# 查找所有超链接
try:
    hyperlink_list = find_hyperlinks(doc_element)
    print(f"通过XML解析找到 {len(hyperlink_list)} 个超链接")
except Exception as e:
    print(f"XML解析时出现错误: {e}")

# 方法2：如果方法1没找到，尝试搜索文档中所有的链接模式
if len(hyperlink_list) == 0:
    print("尝试通过正则表达式查找URL...")
    
    # 收集文档中所有文本
    all_text = ""
    for paragraph in doc.paragraphs:
        all_text += paragraph.text + "\n"
    
    # 查找所有可能的URL模式
    url_patterns = [
        r'https?://[^\s<>"\']+',  # http/https链接
    ]
    
    for pattern in url_patterns:
        matches = re.findall(pattern, all_text, re.IGNORECASE)
        for url in matches:
            # 尝试找到URL周围的文本作为链接文本
            context_start = max(0, all_text.find(url) - 50)
            context_end = min(len(all_text), all_text.find(url) + 50)
            context = all_text[context_start:context_end].strip()
            
            # 清理上下文，提取可能的链接文本
            link_text = context.split('\n')[-1].strip() if '\n' in context else context
            if len(link_text) > 100:  # 如果太长，截断
                link_text = link_text[:100] + "..."
            
            hyperlink_list.append((link_text, url))

# 方法3：查找段落中的超链接（使用python-docx的正确方法）
if len(hyperlink_list) == 0:
    print("尝试通过python-docx的内置方法查找...")
    
    for paragraph in doc.paragraphs:
        # 检查段落中是否有超链接
        if hasattr(paragraph, '_element') and paragraph._element.xpath('.//w:hyperlink'):
            # 提取段落文本作为链接文本
            link_text = paragraph.text.strip()
            
            # 尝试从XML中提取链接
            hyperlinks = paragraph._element.xpath('.//w:hyperlink')
            for hlink in hyperlinks:
                r_id = hlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if r_id and hasattr(doc.part, 'rels'):
                    rel = doc.part.rels.get(r_id)
                    if rel:
                        link_url = rel.target_ref if hasattr(rel, 'target_ref') else str(rel.target)
                        if link_text or link_url:
                            hyperlink_list.append((link_text, link_url))

# 写入结果文件
try:
    with open(txt_file_path, 'w', encoding='utf-8') as f:
        f.write(f"文档 '{os.path.basename(word_file_path)}' 中的超链接列表：\n")
        f.write("=" * 50 + "\n")
        
        if len(hyperlink_list) == 0:
            f.write("未在文档中找到超链接。\n")
            print("注意：未在文档中找到超链接。")
        else:
            # 去重
            unique_links = []
            seen = set()
            for text, url in hyperlink_list:
                # 规范化URL
                norm_url = url.lower().strip()
                if norm_url and norm_url not in seen:
                    seen.add(norm_url)
                    unique_links.append((text, url))
            
            print(f"找到 {len(unique_links)} 个唯一超链接")
            
            for idx, (text, url) in enumerate(unique_links, start=1):
                f.write(f"{text if text else '(无文本)'}")
                f.write(f"    {url}\n")
            
            # 按域名统计
            from urllib.parse import urlparse
            domain_stats = {}
            for _, url in unique_links:
                try:
                    domain = urlparse(url).netloc
                    domain_stats[domain] = domain_stats.get(domain, 0) + 1
                except:
                    pass
            
    
    print(f"结果已保存至：{txt_file_path}")
    
except Exception as e:
    print(f"写入文件时出现错误: {e}")